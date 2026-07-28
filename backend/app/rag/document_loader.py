"""
文档加载器模块
支持多种文档格式的加载和解析
支持格式：.md .pdf .docx .txt

功能：
1. 根据文件扩展名自动选择解析器
2. 文本清洗（去除多余空格、特殊字符）
3. 元数据提取（文件名、路径、类别）
最后将文档内容、类别、元数据封装为Document对象

设计模式：策略模式（不同格式采用不同解析策略）
"""

from typing import List, Optional, Dict, Any
from pathlib import Path
import logging
import re

# 日志记录器
logger = logging.getLogger(__name__)


class Document:
    """
    文档数据类
    统一返回格式，包含文档内容、类别、元数据
    """

    def __init__(self, id: str, content: str, category: str = "", metadata: Optional[Dict[str, Any]] = None):
        self.id = id
        self.content = content
        self.category = category
        self.metadata = metadata if metadata is not None else {}


class DocumentLoader:
    """
    文档加载器
    支持格式：.md .pdf .docx .txt
    
    根据文件扩展名自动选择解析器
    文本清洗（去除多余空格、特殊字符）
    元数据提取（文件名、路径、类别）
    """

    def __init__(self):
        """初始化文档加载器，注册各种格式的解析器"""
        # 解析器注册表：扩展名 -> 解析函数
        self.parsers = {
            ".md": self._parse_markdown,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt": self._parse_txt,
        }

    def load(self, file_path: str, category: str = None) -> Document:
        """
        加载单个文档
        
        参数：
            file_path: 文件路径
            category: 文档类别（可选）
        
        返回：Document对象
        
        异常：
            ValueError: 文件格式不支持
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path) 

        # 检查文件是否存在
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 获取文件扩展名
        ext = path.suffix.lower()

        # 根据扩展名选择解析器
        if ext not in self.parsers:
            raise ValueError(f"不支持的文件格式: {ext}, 支持格式: {list(self.parsers.keys())}")

        # 解析文件
        parser = self.parsers[ext]
        content = parser(path)

        # 文本清洗
        content = self._clean_text(content)

        # 提取元数据
        metadata = self._extract_metadata(path, category)

        # 生成文档ID（基于文件路径的哈希）
        doc_id = str(hash(str(path)))

        logger.info(f"文档加载完成: {path.name}, 内容长度: {len(content)}")

        return Document(
            id=doc_id,
            content=content,
            category=category or "",
            metadata=metadata
        )

    def load_directory(self, dir_path: str, category: str = None) -> List[Document]:
        """
        加载目录下所有支持格式的文档
        
        参数：
            dir_path: 目录路径
            category: 文档类别（可选）
        
        返回：Document列表
        
        异常：
            FileNotFoundError: 目录不存在
        """
        directory = Path(dir_path)

        # 检查目录是否存在
        if not directory.exists():
            raise FileNotFoundError(f"目录不存在: {dir_path}")

        # 收集所有支持的文件
        all_files = []
        for ext in self.parsers.keys():
            all_files.extend(directory.glob(f"**/*{ext}"))

        # 加载每个文件
        documents = []
        for file_path in all_files:
            try:
                doc = self.load(str(file_path), category)
                documents.append(doc)
            except Exception as e:
                logger.warning(f"加载文件失败: {file_path}, 错误: {e}")

        logger.info(f"目录加载完成: {dir_path}, 共加载{len(documents)}个文档")

        return documents

    def _parse_markdown(self, path: Path) -> str:
        """
        解析Markdown文件
        
        参数：
            path: 文件路径
        
        返回：
            提取的文本内容
        """
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        # 简单处理：去除Markdown格式，保留纯文本
        # 移除标题标记
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        # 移除粗体/斜体标记
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        # 移除链接标记
        content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
        # 移除代码块标记
        content = re.sub(r'```[\s\S]*?```', '', content)
        # 移除行内代码标记
        content = re.sub(r'`([^`]+)`', r'\1', content)
        # 移除图片标记
        content = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', content)

        return content

    def _parse_pdf(self, path: Path) -> str:
        """
        解析PDF文件
        
        参数：path: 文件路径
        
        返回：提取的文本内容
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("需要安装 PyPDF2 库: pip install PyPDF2")

        reader = PdfReader(str(path))
        content = ""

        for page in reader.pages:
            text = page.extract_text()
            if text:
                content += text + "\n"

        return content.strip()

    def _parse_docx(self, path: Path) -> str:
        """
        解析Word文件(.docx)
        
        参数：
            path: 文件路径
        
        返回：
            提取的文本内容
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")

        doc = DocxDocument(str(path))
        content = ""

        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_content = "\t".join(cell.text for cell in row.cells)
                content += row_content + "\n"

        return content.strip()

    def _parse_txt(self, path: Path) -> str:
        """
        解析文本文件
        
        参数：
            path: 文件路径
        
        返回：
            提取的文本内容
        """
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        return content

    def _clean_text(self, text: str) -> str:
        """
        文本清洗
        
        处理内容：
        去除多余空格和换行
        去除特殊字符
        统一换行符
        
        参数：text: 原始文本
        
        返回：清洗后的文本
        """
        if not text:
            return ""

        # 去除首尾空白
        text = text.strip()

        # 统一换行符为\n
        text = text.replace("\r\n", "\n")
        text = text.replace("\r", "\n")

        # 去除连续的空行（保留最多两个换行）
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 去除每行首尾的多余空格
        lines = text.split("\n")
        lines = [line.strip() for line in lines]
        text = "\n".join(lines)

        # 去除连续的空格（保留一个）
        text = re.sub(r' {2,}', ' ', text)

        # 去除特殊控制字符（保留基本的空白字符）
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        return text.strip()

    def _extract_metadata(self, path: Path, category: str = None) -> Dict[str, Any]:
        """
        提取元数据
        参数：
            path: 文件路径
            category: 文档类别
        
        返回：元数据字典
        """
        metadata = {
            "file_name": path.name,
            "file_path": str(path),
            "file_extension": path.suffix,
            "category": category or "",
        }

        # 获取文件大小
        try:
            metadata["file_size"] = path.stat().st_size
        except Exception as e:
            logger.warning(f"获取文件大小失败: {e}")

        # 获取文件修改时间
        try:
            metadata["modified_time"] = path.stat().st_mtime
        except Exception as e:
            logger.warning(f"获取文件修改时间失败: {e}")

        return metadata


# ======================== 文件内自测脚本 ========================
# 运行方式：在 backend/app 目录下执行  python -m rag.document_loader
if __name__ == "__main__":
    print("=" * 60)
    print("文档加载器自测开始")
    print("=" * 60)

    # 创建文档加载器实例
    loader = DocumentLoader()

    # ---------- 测试1：文本清洗功能 ----------
    # 用例：清洗多余空格、换行、特殊字符
    dirty_text = """
        这是一段   测试文本，包含多余的   空格。
        
        还有多个换行。
        
        以及一些特殊字符：\t\n\x00
        """
    cleaned = loader._clean_text(dirty_text)
    assert "   " not in cleaned, "多余空格未被清理"
    assert "\n\n\n" not in cleaned, "多余换行未被清理"
    assert "\x00" not in cleaned, "特殊字符未被清理"
    print("[通过] 测试1 - 文本清洗功能")

    # ---------- 测试2：元数据提取 ----------
    # 用例：从文件路径提取元数据
    import tempfile
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("测试内容")
        temp_path = f.name

    metadata = loader._extract_metadata(Path(temp_path), category="测试")
    assert metadata["file_name"] == Path(temp_path).name
    assert metadata["file_path"] == temp_path
    assert metadata["file_extension"] == ".txt"
    assert metadata["category"] == "测试"
    assert "file_size" in metadata
    assert "modified_time" in metadata
    import os
    os.unlink(temp_path)
    print("[通过] 测试2 - 元数据提取")

    # ---------- 测试3：加载TXT文件 ----------
    # 用例：创建临时TXT文件并加载
    with tempfile.NamedTemporaryFile(mode="w", encoding="utf-8", suffix=".txt", delete=False) as f:
        f.write("苹果富含维生素C，每天吃一个有益健康。\n香蕉含有丰富的钾元素。")
        txt_path = f.name

    doc = loader.load(txt_path, category="营养学")
    assert doc.id is not None
    assert "苹果" in doc.content
    assert "香蕉" in doc.content
    assert doc.category == "营养学"
    assert doc.metadata["file_name"] == Path(txt_path).name
    os.unlink(txt_path)
    print("[通过] 测试3 - 加载TXT文件")

    # ---------- 测试4：加载Markdown文件 ----------
    # 用例：创建临时Markdown文件并加载，验证格式被清除
    with tempfile.NamedTemporaryFile(mode="w", encoding="utf-8", suffix=".md", delete=False) as f:
        f.write("""# 水果营养
        **苹果**富含维生素C，[了解更多](https://example.com)。
        
        ```python
        print("代码块")
        ```
        """)
        md_path = f.name

    doc = loader.load(md_path, category="营养学")
    assert "#" not in doc.content, "标题标记未被清除"
    assert "**" not in doc.content, "粗体标记未被清除"
    assert "[" not in doc.content or "](http" not in doc.content, "链接标记未被清除"
    assert "```" not in doc.content, "代码块标记未被清除"
    assert "苹果" in doc.content
    os.unlink(md_path)
    print("[通过] 测试4 - 加载Markdown文件")

    # ---------- 测试5：不支持的格式 ----------
    # 用例：加载不支持的格式，验证抛出异常
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
        f.write("test")
        csv_path = f.name

    try:
        loader.load(csv_path)
        assert False, "应该抛出不支持格式异常"
    except ValueError as e:
        assert "不支持的文件格式" in str(e)
    os.unlink(csv_path)
    print("[通过] 测试5 - 不支持格式异常处理")

    # ---------- 测试6：不存在的文件 ----------
    # 用例：加载不存在的文件，验证抛出异常
    try:
        loader.load("/path/to/nonexistent/file.txt")
        assert False, "应该抛出文件不存在异常"
    except FileNotFoundError:
        pass
    print("[通过] 测试6 - 不存在文件异常处理")

    print("=" * 60)
    print("文档加载器自测全部通过（6/6）")
    print("=" * 60)